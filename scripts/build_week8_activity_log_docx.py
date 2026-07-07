#!/usr/bin/env python3
"""Build Week 8 IST495 activity log docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT = PROJECT_ROOT / "reports" / "weekly_updates" / "week8_activity_log_IST495.docx"

ROWS = [
    (
        "Monday",
        "Morning: 9:30 – 12:00",
        "Looked up how Stocktwits API works (professor roadmap social part). Read their JSON for messages and dates.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Monday",
        "Afternoon: 1:30 – 4:00",
        "Wrote collect_stocktwits.py to pull messages for each ticker from the public API.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Tuesday",
        "Morning: 10:00 – 12:30",
        "Made live_stocktwits_metrics.py — fetch 20 tickers, count posts in same date range as Finviz.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Tuesday",
        "Afternoon: 2:00 – 4:30",
        "Added Stocktwits to dashboard: new tab, two table columns, updated live fetch line.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Wednesday",
        "Morning: 9:00 – 11:30",
        "Fixed date parsing for Stocktwits. Unit tests with fake JSON.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Wednesday",
        "Afternoon: 1:00 – 3:30",
        "Tested on Mac. Finviz OK (15 news / 7d). Stocktwits curl gives HTML on my wifi — not token issue.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Thursday",
        "Morning: 10:00 – 12:00",
        "Fixed crash when Stocktwits empty (KeyError).",
        "Individual",
        "2 hrs",
    ),
    (
        "Thursday",
        "Afternoon: 2:00 – 4:00",
        "Shortened yellow/blue warning text on dashboard.",
        "Individual",
        "2 hrs",
    ),
    (
        "Friday",
        "Morning: 9:30 – 11:30",
        "Practiced demo script for professor (Week 7 live + Week 8 Stocktwits).",
        "Individual",
        "2 hrs",
    ),
    (
        "Friday",
        "Afternoon: 1:00 – 3:00",
        "Wrote week8 notes and knowledge pack. Ran pytest.",
        "Individual",
        "2 hrs",
    ),
    (
        "Saturday",
        "Morning: 10:00 – 12:00",
        "Read Canvas video intern posts. Mockup is separate from weekly log.",
        "Individual",
        "2 hrs",
    ),
    (
        "Sunday",
        "Afternoon: 2:00 – 4:30",
        "Ran Streamlit again (4 tabs). Finished activity log.",
        "Individual",
        "2.5 hrs",
    ),
]

COMMENTS = (
    "I liked seeing Finviz and Stocktwits in the same dashboard. Finviz works with my token but "
    "Stocktwits on my home wifi only gives HTML back — took me a while to see that's not the same "
    "as a wrong password. social_density next to message_density helps compare news vs social for the same week."
)

EXTERNAL_HELP = (
    "Cursor / AI helper (~4–5 hrs): Stocktwits code, dashboard, crash fix, alert wording, log draft. "
    "Stocktwits and Streamlit docs (~1 hr). I still tested Finviz live on my own laptop."
)

LINKS = (
    "https://elite.finviz.com/api_explanation\n"
    "https://api.stocktwits.com/developers/docs\n"
    "https://docs.streamlit.io/\n"
    "https://feedflash-production.up.railway.app/\n"
    "https://github.com/jml8284/fin-news-sentiment"
)

CONTRIBUTIONS = (
    "1. Stocktwits fetch code (collect_stocktwits.py).\n"
    "2. Post counts + social_density (live_stocktwits_metrics.py).\n"
    "3. Dashboard Stocktwits tab and table columns.\n"
    "4. Empty Stocktwits crash fix; simpler warnings.\n"
    "5. tests/test_stocktwits.py.\n"
    "6. Week 7 Finviz live still works; Stocktwits in code even if my network shows zero."
)


def main() -> None:
    doc = Document()
    doc.add_paragraph("Name: Jinyang Liu")
    doc.add_paragraph("Email: jml8284@psu.edu")
    doc.add_paragraph(
        "Note: Please complete all columns, specially the last two columns. Thank You."
    )
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [
        "Day of week",
        "Time of Day",
        "From - To",
        "Description of Activity",
        "Individual or Group?",
        "Duration",
    ]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text

    for row in ROWS:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    doc.add_paragraph("")
    doc.add_paragraph(
        "Comments: Please comment a learning experience you enjoyed this week."
    )
    doc.add_paragraph(COMMENTS)
    doc.add_paragraph("")
    doc.add_paragraph(
        "External Help: Please outline any help you have received from outside resources "
        "such as ChatGPT or tutors. If so, in what areas and for how long?"
    )
    doc.add_paragraph(EXTERNAL_HELP)
    doc.add_paragraph("")
    doc.add_paragraph(
        "Please list the link of any external materials you have used to assist you "
        "with your course project. This could be Youtube link, LinkedIn links, etc."
    )
    doc.add_paragraph(LINKS)
    doc.add_paragraph("")
    doc.add_paragraph("What was your contributions to the course project?")
    doc.add_paragraph(CONTRIBUTIONS)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
