#!/usr/bin/env python3
"""Build Week 7 IST495 activity log docx from project content."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT = PROJECT_ROOT / "reports" / "weekly_updates" / "week7_activity_log_IST495.docx"

ROWS = [
    (
        "Monday",
        "Morning: 9:00 – 12:00",
        "Researched FinBERT (ProsusAI/finbert) for professor roadmap item 6 (AI rankings). "
        "Started src/sentiment_engines.py with pluggable VADER / FinBERT engines.",
        "Individual",
        "3 hrs",
    ),
    (
        "Monday",
        "Afternoon: 1:30 – 3:30",
        "Updated sentiment_analysis.py with --engine vader|finbert; added sentiment_engine.txt metadata. "
        "Installed transformers and torch.",
        "Individual",
        "2 hrs",
    ),
    (
        "Tuesday",
        "Morning: 10:00 – 12:30",
        "Extended evaluate_sentiment.py to compare VADER and FinBERT on PhraseBank and combined benchmarks. "
        "Generated sentiment_eval_report.csv.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Tuesday",
        "Afternoon: 2:00 – 4:30",
        "Wired run_pipeline.py --finbert and --evaluate-models. Added tests/test_sentiment_engines.py. "
        "First FinBERT benchmark (~76% vs ~57% on PhraseBank).",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Wednesday",
        "Morning: 9:30 – 12:00",
        "Created src/live_finviz_metrics.py for parallel live Finviz quote-page news scrape (not CSV snapshots).",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Wednesday",
        "Afternoon: 1:00 – 3:30",
        "Rewrote dashboard ranked table and news viewer for live Finviz data; live fetch timestamp; "
        "60s screener/chart refresh; Refresh now button.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Thursday",
        "Morning: 10:00 – 12:30",
        "Fixed Finviz news HTML parser; removed 80-item cap; elite URL fallback; trust_env=False for proxy issues.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Thursday",
        "Afternoon: 2:00 – 4:00",
        "Added live metrics tests; ran scripts/verify_live_finviz.py; confirmed live screener + quote news with Elite token.",
        "Individual",
        "2 hrs",
    ),
    (
        "Friday",
        "Morning: 9:00 – 11:30",
        "Fixed dashboard long loading: st.fragment was re-triggering full news scrape every 60s. "
        "Separated 5-min news cache from 60s screener/chart refresh.",
        "Individual",
        "2.5 hrs",
    ),
    (
        "Friday",
        "Afternoon: 1:00 – 3:00",
        "Live Dashboard uses VADER for speed; FinBERT for offline pipeline and benchmark. "
        "Fixed Week 6 chart change % stale outside fragment.",
        "Individual",
        "2 hrs",
    ),
    (
        "Saturday",
        "Morning: 10:00 – 12:00",
        "Auto-generated sentiment_eval_report.md; updated message_density to Sparse / Moderate / Dense in UI.",
        "Individual",
        "2 hrs",
    ),
    (
        "Saturday",
        "Afternoon: 1:30 – 3:00",
        "35 pytest tests passing; updated week7_update.md; rehearsed professor demo (live fetch + FinBERT report).",
        "Individual",
        "1.5 hrs",
    ),
    (
        "Sunday",
        "Afternoon: 2:00 – 4:00",
        "Drafted Week 7 activity log; verified Streamlit dashboard end-to-end with live Finviz Elite.",
        "Individual",
        "2 hrs",
    ),
]

COMMENTS = (
    "The biggest breakthrough was moving from pipeline CSV snapshots to true live Finviz scraping in the dashboard. "
    "Seeing the Live fetch UTC timestamp update and news_count come directly from quote pages made the project feel "
    "like a real screener. Integrating FinBERT also clarified the speed vs accuracy trade-off: VADER for live scoring, "
    "FinBERT for benchmark evaluation (professor roadmap item 6 — AI rankings)."
)

EXTERNAL_HELP = (
    "Cursor / AI coding assistant (~5–6 hrs): FinBERT engine module, live_finviz_metrics refactor, dashboard "
    "fragment/cache debugging, Finviz parser fixes, pytest, activity log draft. "
    "Hugging Face FinBERT docs (~45 min). Streamlit docs (~30 min). Finviz Elite API (~30 min). "
    "I reviewed and tested all changes myself."
)

LINKS = (
    "• Finviz Elite API: https://elite.finviz.com/api_explanation\n"
    "• Finviz screener: https://elite.finviz.com/screener\n"
    "• ProsusAI FinBERT: https://huggingface.co/ProsusAI/finbert\n"
    "• Streamlit: https://docs.streamlit.io/\n"
    "• VADER: https://github.com/cjhutto/vaderSentiment\n"
    "• FeedFlash reference: https://feedflash-production.up.railway.app/\n"
    "• Project repo: https://github.com/jml8284/fin-news-sentiment"
)

CONTRIBUTIONS = (
    "1. FinBERT integration (sentiment_engines.py): pluggable VADER/FinBERT; pipeline --engine finbert.\n"
    "2. Model evaluation: VADER vs FinBERT on PhraseBank + combined; sentiment_eval_report.csv and .md.\n"
    "3. Live Finviz dashboard: live_finviz_metrics.py + dashboard rewrite; live fetch timestamp; 60s refresh.\n"
    "4. Finviz news reliability: collect_news.py parser fixes, no row cap, elite fallback.\n"
    "5. Performance/UX: fixed fragment re-scrape bug; VADER live scoring; Sparse/Moderate/Dense density labels.\n"
    "6. Testing/docs: 35 pytest tests; week7_update.md; verify_live_finviz.py."
)


def main() -> None:
    doc = Document()
    doc.add_paragraph("Name: Jinyang Liu")
    doc.add_paragraph("Email: jml8284@psu.edu")
    doc.add_paragraph(
        "Note: Please complete all columns, specially the last two columns. Thank You."
    )
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [
        "Day of week",
        "Time of Day",
        "From - To",
        "Description of Activity",
        "Individual or Group?",
        "Duration",
    ]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text

    for row in ROWS:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    doc.add_paragraph("")
    doc.add_paragraph(
        "Comments: Please comment a learning experience you enjoyed this week."
    )
    doc.add_paragraph(COMMENTS)
    doc.add_paragraph("")
    doc.add_paragraph(
        "External Help: Please outline any help you have received from outside resources "
        "such as ChatGPT or tutors. If so, in what areas and for how long?"
    )
    doc.add_paragraph(EXTERNAL_HELP)
    doc.add_paragraph("")
    doc.add_paragraph(
        "Please list the link of any external materials you have used to assist you "
        "with your course project. This could be Youtube link, LinkedIn links, etc."
    )
    doc.add_paragraph(LINKS)
    doc.add_paragraph("")
    doc.add_paragraph("What was your contributions to the course project?")
    doc.add_paragraph(CONTRIBUTIONS)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
